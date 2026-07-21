"""Word path — extract products directly from a structured .docx spec/
estimate table (no OCR / table-detection needed: the data is already in
table cells).

Mirrors excel_extractor.py's approach and quirks, since GOST/smeta-style
estimates are exported to both .xlsx and .docx with the same column layout
and the same junk rows:
  - phantom "legend" rows where cells just contain their own column position
    number (1, 2, 3, ... 7) instead of real data.
  - a code cell holding two lines: the article code and a price-reference /
    norm line (e.g. "522-101-1212008\\nСпрСЦ 05.2024") — only the first line
    is a real code, used here just to validate the row is a real product.
  - trailing fully-empty rows at the end of the table.

Unlike Excel, a .docx has no merged-cell "anchor" concept exposed the same
way: python-docx repeats the same Cell object for every grid position a
merge spans, so reading the name column directly already gives the full
(merged) text with no extra handling needed.

A .docx can contain more than one table (e.g. a cover/title table before
the actual estimate). By default we scan every table and use whichever one
yields product rows; pass `table_index` to pin a specific one instead.

Column positions default to the smeta/GOST layout seen in the .xlsx sample
(same as ExcelColumnMap):
    № (1) | Шифр+справочник (2) | Наименование (3) |
    ... | Ед.изм (7) | Кол-во (8) | Цена (9) | Сумма (10)
Pass a different WordColumnMap for templates that don't match this layout.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Optional

from docx import Document

from .config import ParserConfig
from .models import Product
from .text_utils import normalize_ws, normalize_name, map_headers, parse_quantity

# Same stop words as excel_extractor.py / product_extractor.py's PDF path.
_STOP_WORDS = (
    "итого", "всего", "total", "subtotal", "sum", "сумма прописью",
    "раздел", "section",
)

# A real article code looks like "522-101-1212008" / "245-401-0201".
_CODE_RE = re.compile(r"\d{2,3}-\d{2,3}-\d+")


@dataclass
class WordColumnMap:
    """1-based column (cell) indices within a table row."""
    item_no: int = 1
    code: int = 2
    name: int = 3
    unit: int = 7
    quantity: int = 8
    price: int = 9
    total: int = 10


class WordProductExtractor:
    def __init__(
        self,
        column_map: Optional[WordColumnMap] = None,
        table_index: Optional[int] = None,
        parser_config: Optional[ParserConfig] = None,
    ):
        self.column_map = column_map
        self.table_index = table_index
        self.parser_config = parser_config or ParserConfig()

    def extract(self, docx_path: str) -> list[Product]:
        doc = Document(docx_path)
        if not doc.tables:
            return []

        tables = (
            [doc.tables[self.table_index]]
            if self.table_index is not None
            else doc.tables
        )

        products: list[Product] = []
        for table in tables:
            products.extend(self._extract_table(table))

        return self._dedupe(products)

    def _extract_table(self, table) -> list[Product]:
        rows = []
        for row in table.rows:
            cells = [normalize_ws(cell.text) for cell in row.cells]
            if any(cells):
                rows.append(cells)

        if not rows:
            return []

        if self.column_map is not None:
            header_map = {
                "item_no": self.column_map.item_no - 1,
                "code": self.column_map.code - 1,
                "name": self.column_map.name - 1,
                "unit": self.column_map.unit - 1,
                "quantity": self.column_map.quantity - 1,
                "price": self.column_map.price - 1,
                "total": self.column_map.total - 1,
            }
            header_row_index = 0
        else:
            header_row_index, header_map = self._find_best_header(rows)
            if header_map is None:
                return []

        cm = WordColumnMap(
            item_no=header_map.get("item_no", 0) + 1,
            code=header_map.get("code", 1) + 1,
            name=header_map.get("name", 2) + 1,
            unit=header_map.get("unit", 6) + 1,
            quantity=header_map.get("quantity", 7) + 1,
            price=header_map.get("price", 8) + 1,
            total=header_map.get("total", 9) + 1,
        )

        products: list[Product] = []
        for row in rows[header_row_index + 1 :]:
            code_cell = row[cm.code - 1] if cm.code - 1 < len(row) else ""
            name_cell = row[cm.name - 1] if cm.name - 1 < len(row) else ""
            qty_cell = row[cm.quantity - 1] if cm.quantity - 1 < len(row) else ""

            if not name_cell and not code_cell:
                continue
            if name_cell.isdigit() and code_cell.isdigit():
                continue

            joined = (name_cell + " " + code_cell).lower()
            if any(sw in joined for sw in _STOP_WORDS):
                continue

            name = normalize_name(name_cell)
            if not name:
                continue

            quantity = parse_quantity(qty_cell)
            products.append(Product(name=name, quantity=quantity))

        return products

    def _find_best_header(self, rows: list[list[str]]) -> tuple[int, dict[str, int] | None]:
        best_map: dict[str, int] | None = None
        best_score = -1
        best_idx = 0
        for idx, row in enumerate(rows[:8]):
            if not any(row):
                continue
            cmap = map_headers(row, self.parser_config)
            if not cmap:
                continue
            score = len(cmap) + (1 if "name" in cmap or "description" in cmap else 0)
            if score > best_score:
                best_score = score
                best_map = cmap
                best_idx = idx
        return best_idx, best_map

    @staticmethod
    def _to_number(value) -> Optional[float]:
        if value is None:
            return None
        s = normalize_ws(str(value)).replace(" ", "").replace(",", ".")
        if not s:
            return None
        try:
            return float(s)
        except ValueError:
            return None

    @staticmethod
    def _dedupe(products: list[Product]) -> list[Product]:
        seen: set[tuple] = set()
        out: list[Product] = []
        for p in products:
            key = (p.name.lower(), p.quantity)
            if key in seen:
                continue
            seen.add(key)
            out.append(p)
        return out
